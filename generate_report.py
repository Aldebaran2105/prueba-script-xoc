"""Genera DOCX de soporte desde las plantillas locales de XOC.

No modifica las plantillas. Azure redacta las secciones narrativas usando solo el JSON
de entrada; el modo --no-azure existe únicamente para probar el armado del documento.
"""

from __future__ import annotations

import argparse
import io
import json
import os
import re
import shutil
import zipfile
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from dotenv import load_dotenv

ROOT = Path(__file__).resolve().parent
TEMPLATES = {
    "small": ROOT / "plantillas" / "Plantilla Small Report.docx",
    "soporte": ROOT / "plantillas" / "Plantilla Informe Soporte.docx",
}
OUTPUT_DIR = ROOT / "output"
BRAND_BLUE = RGBColor(0x00, 0x6F, 0xA1)
TEXT_GRAY = RGBColor(0x59, 0x59, 0x59)
FOOTER_GRAY = RGBColor(0x76, 0x71, 0x71)
TABLE_HEADER = "006FA1"
TABLE_ALT = "EAF4F8"
CALLOUT_BLUE = "EAF4F8"
CALLOUT_YELLOW = "FFF7DE"
CALLOUT_GRAY = "F4F7F9"
BORDER_LIGHT = "C9DCE8"


def _fit_cover_font(text: str, max_width: int, start_size: int = 27, min_size: int = 15) -> Any:
    from PIL import Image, ImageDraw, ImageFont

    probe = Image.new("RGB", (max_width, 80), "white")
    draw = ImageDraw.Draw(probe)
    for size in range(start_size, min_size - 1, -1):
        font = ImageFont.truetype("arialbd.ttf", size)
        if draw.textbbox((0, 0), text, font=font)[2] <= max_width:
            return font
    return ImageFont.truetype("arialbd.ttf", min_size)


def patch_small_cover_branding(docx_path: Path, client_name: str) -> None:
    """Actualiza textos de cliente quemados en la imagen de portada small."""
    client = str(client_name or "Cliente").strip()[:42]
    tmp_path = docx_path.with_suffix(".tmp.docx")
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception as exc:
        raise RuntimeError("Pillow es requerido para actualizar la portada small") from exc

    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == "word/media/image1.png":
                image = Image.open(io.BytesIO(data)).convert("RGBA")
                draw = ImageDraw.Draw(image)

                # Logo/nombre de la plantilla en la esquina superior izquierda.
                draw.rounded_rectangle((34, 92, 184, 282), radius=8, fill=(7, 25, 40, 225))
                top_font = _fit_cover_font(client, 120, start_size=16, min_size=10)
                words = client.split()
                lines: list[str] = []
                current = ""
                for word in words:
                    candidate = f"{current} {word}".strip()
                    if len(candidate) <= 17:
                        current = candidate
                    else:
                        if current:
                            lines.append(current)
                        current = word
                if current:
                    lines.append(current)
                wrapped = "\n".join(lines[:4]) or client
                draw.multiline_text((48, 164), wrapped, fill=(255, 255, 255, 255), font=top_font, spacing=4, align="center")

                # Valor de Prepared for.
                draw.rectangle((78, 842, 380, 900), fill=(255, 255, 255, 255))
                label_font = ImageFont.truetype("arial.ttf", 21)
                draw.text((84, 844), "Prepared for :", fill=(9, 32, 70, 255), font=label_font)
                font = _fit_cover_font(client, 282, start_size=25, min_size=16)
                draw.text((84, 870), client, fill=(9, 32, 70, 255), font=font)

                buffer = io.BytesIO()
                image.save(buffer, format="PNG")
                data = buffer.getvalue()
            target.writestr(item, data)
    shutil.move(str(tmp_path), str(docx_path))


def patch_support_cover_text_color(docx_path: Path, values: list[str], color: str = "000000") -> None:
    """Fuerza color legible solo en los valores dinamicos de portada soporte."""
    from lxml import etree

    wanted = {str(value).strip() for value in values if str(value).strip()}
    tmp_path = docx_path.with_suffix(".tmp.docx")
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == "word/document.xml":
                root = etree.fromstring(data)
                for box in root.xpath(".//w:txbxContent", namespaces=namespace):
                    text = "".join(box.xpath(".//w:t/text()", namespaces=namespace)).strip()
                    if text not in wanted:
                        continue
                    for color_node in box.xpath(".//w:color", namespaces=namespace):
                        color_node.set(qn("w:val"), color)
                        for attr in (qn("w:themeColor"), qn("w:themeShade"), qn("w:themeTint")):
                            if attr in color_node.attrib:
                                del color_node.attrib[attr]
                    for run in box.xpath(".//w:r", namespaces=namespace):
                        run_properties = run.find(qn("w:rPr"))
                        if run_properties is None:
                            run_properties = OxmlElement("w:rPr")
                            run.insert(0, run_properties)
                        color_node = run_properties.find(qn("w:color"))
                        if color_node is None:
                            color_node = OxmlElement("w:color")
                            run_properties.append(color_node)
                        color_node.set(qn("w:val"), color)
                        for attr in (qn("w:themeColor"), qn("w:themeShade"), qn("w:themeTint")):
                            if attr in color_node.attrib:
                                del color_node.attrib[attr]
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            target.writestr(item, data)
    shutil.move(str(tmp_path), str(docx_path))


def _safe_name(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9]+", "-", value.upper().strip())
    return value.strip("-")[:70] or "CLIENTE"


def _report_date(data: dict[str, Any]) -> str:
    return str(data.get("report_date") or date.today().isoformat())


def _report_date_spanish(data: dict[str, Any]) -> str:
    raw_date = _report_date(data)
    try:
        parsed = date.fromisoformat(raw_date)
    except ValueError:
        return raw_date
    months = ("enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre")
    return f"{parsed.day} de {months[parsed.month - 1]} de {parsed.year}"


def output_path(kind: str, data: dict[str, Any]) -> Path:
    prefix = "SMALL-REPORT" if kind == "small" else "INFORME-SOPORTE"
    case = _safe_name(str(data.get("case_number") or "SIN-CASO"))
    client = _safe_name(str(data.get("client_name") or data.get("establishment") or "CLIENTE"))
    return OUTPUT_DIR / f"{prefix}_{case}_{client}_{_report_date(data)}.docx"


def _as_list(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item) for item in value if str(item).strip()]
    return [str(value)] if value else []


def local_sections(kind: str, data: dict[str, Any]) -> dict[str, str]:
    actions = "; ".join(_as_list(data.get("actions_performed"))) or "No se registraron acciones."
    results = "; ".join(_as_list(data.get("results"))) or "No se registraron resultados."
    recommendations = "; ".join(_as_list(data.get("recommendations"))) or "No se registraron recomendaciones."
    if kind == "small":
        return {
            "information_gathering": str(data.get("information_gathering") or data.get("request_description") or "Sin información adicional."),
            "possible_causes": "; ".join(_as_list(data.get("possible_causes"))) or "Sin hipótesis registradas.",
            "actions": actions,
            "results": results,
            "recommendations": recommendations,
        }
    return {
        "background": str(data.get("background") or "Sin antecedentes adicionales."),
        "initial_scenario": str(data.get("initial_scenario") or "Sin escenario inicial adicional."),
        "actions": actions,
        "results": results,
        "recommendations": recommendations,
        "conclusion": f"La actividad relacionada con {data.get('subject', 'el soporte')} fue documentada. {results}",
    }


def _prompt(kind: str, data: dict[str, Any]) -> str:
    fields = (
        "information_gathering, possible_causes, actions, results, recommendations"
        if kind == "small"
        else "background, initial_scenario, actions, results, recommendations, conclusion"
    )
    return (
        "Eres un analista senior de soporte TI de XOC. Redacta un informe técnico en español. "
        "Usa exclusivamente los datos JSON recibidos: no inventes fechas, configuraciones, resultados, activos ni acciones. "
        "Mantén un tono profesional, claro y verificable. Devuelve SOLO un JSON válido sin Markdown, "
        f"con exactamente estas claves: {fields}. Cada valor debe ser un texto conciso.\n\n"
        f"Datos: {json.dumps(data, ensure_ascii=False, separators=(',', ':'))}"
    )


def _parse_json(raw: str, expected: set[str]) -> dict[str, str]:
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=re.IGNORECASE)
    parsed = json.loads(raw)
    if not isinstance(parsed, dict) or set(parsed) != expected:
        raise RuntimeError("Azure devolvió una estructura distinta a la solicitada")
    return {key: str(value).strip() for key, value in parsed.items()}


def generate_with_azure(kind: str, data: dict[str, Any]) -> dict[str, str]:
    expected = (
        {"information_gathering", "possible_causes", "actions", "results", "recommendations"}
        if kind == "small"
        else {"background", "initial_scenario", "actions", "results", "recommendations", "conclusion"}
    )
    prompt = _prompt(kind, data)
    api_key = os.environ.get("AZURE_OPENAI_API_KEY", "").strip()
    if api_key:
        from openai import OpenAI

        endpoint = os.environ["AZURE_OPENAI_ENDPOINT"].strip().rstrip("/")
        deployment = os.environ["AZURE_OPENAI_DEPLOYMENT"].strip()
        response = OpenAI(base_url=f"{endpoint}/openai/v1", api_key=api_key).responses.create(
            model=deployment, input=prompt, max_output_tokens=1600, reasoning={"effort": "low"}
        )
        return _parse_json(response.output_text, expected)

    from azure.ai.projects import AIProjectClient
    from azure.identity import DefaultAzureCredential

    endpoint = os.environ["AZURE_FOUNDRY_PROJECT_ENDPOINT"].strip()
    agent_name = os.environ["AZURE_FOUNDRY_AGENT_NAME"].strip()
    client = AIProjectClient(endpoint=endpoint, credential=DefaultAzureCredential())
    agents = client.agents
    agent = agents.get(agent_name) if hasattr(agents, "get") else agents.get_agent(agent_name)
    thread = agents.threads.create()
    agents.messages.create(thread_id=thread.id, role="user", content=prompt)
    run = agents.runs.create_and_process(thread_id=thread.id, agent_id=agent.id)
    if str(getattr(run, "status", "")).lower() != "completed":
        raise RuntimeError(f"Azure Foundry no completó la ejecución: {getattr(run, 'last_error', None)}")
    messages = list(agents.messages.list(thread_id=thread.id))
    for message in messages:
        if getattr(message, "role", "") != "assistant":
            continue
        content = getattr(message, "content", []) or []
        text = "\n".join(item.text.value for item in content if getattr(item, "text", None) and hasattr(item.text, "value"))
        if text:
            return _parse_json(text, expected)
    raise RuntimeError("Azure Foundry no devolvió contenido del agente")


def clear_document_body_except_cover(document: Document) -> None:
    """Conserva la portada gráfica y su salto de sección; elimina solo el ejemplo.

    Las plantillas no tienen placeholders: su portada está formada por imágenes y
    elementos de dibujo en el primer párrafo. El siguiente salto de sección define
    dónde debe empezar el informe real.
    """
    body = document._element.body
    children = list(body)
    first_paragraph_section_break = next((index for index, child in enumerate(children) if child.xpath("./w:pPr/w:sectPr")), None)
    if first_paragraph_section_break is not None:
        preserved = set(children[: first_paragraph_section_break + 1])
    else:
        preserved = {
            child
            for child in children
            if child.tag != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"
        }
    for child in list(body):
        if child not in preserved and child.tag != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr":
            body.remove(child)


def _textbox_text(box: Any) -> str:
    return "".join(node.text or "" for node in box.iter(qn("w:t")))


def _textbox_styles(box: Any) -> str:
    styles: list[str] = []
    node = box
    while node is not None:
        style = node.get("style")
        if style:
            styles.append(style)
        node = node.getparent()
    return " ".join(styles)


def _set_textbox_text(box: Any, value: str, color: str | None = None) -> None:
    nodes = list(box.iter(qn("w:t")))
    if not nodes:
        return
    nodes[0].text = value
    for node in nodes[1:]:
        node.text = ""
    if color:
        _set_textbox_color(box, color)


def _set_textbox_color(box: Any, color: str) -> None:
    for color_node in box.iter(qn("w:color")):
        color_node.set(qn("w:val"), color)
        for attr in (qn("w:themeColor"), qn("w:themeShade"), qn("w:themeTint")):
            if attr in color_node.attrib:
                del color_node.attrib[attr]
    for run in box.iter(qn("w:r")):
        run_properties = run.find(qn("w:rPr"))
        if run_properties is None:
            run_properties = OxmlElement("w:rPr")
            run.insert(0, run_properties)
        color_node = run_properties.find(qn("w:color"))
        if color_node is None:
            color_node = OxmlElement("w:color")
            run_properties.append(color_node)
        color_node.set(qn("w:val"), color)
        for attr in (qn("w:themeColor"), qn("w:themeShade"), qn("w:themeTint")):
            if attr in color_node.attrib:
                del color_node.attrib[attr]


def update_cover_fields(document: Document, kind: str, data: dict[str, Any]) -> None:
    """Reemplaza los cuadros de texto existentes de la portada sin alterar su diseño."""
    client = str(data.get("client_name") or data.get("establishment") or "Cliente")[:55]
    subject = str(data.get("subject") or "Informe de soporte")[:75]
    report_date = _report_date_spanish(data)
    prepared_by = str(data.get("prepared_by") or "TxDxSecure")[:40]
    title = str(data.get("report_title") or "Informe de Soporte")[:50]

    change_boxes: list[Any] = []
    for box in document._element.xpath(".//w:txbxContent"):
        original = _textbox_text(box).strip()
        if kind == "small" and original == "Change Tenant":
            _set_textbox_text(box, client)
        elif kind == "small" and original.startswith("REPORTE"):
            _set_textbox_text(box, f"SMALL REPORT — {client} — CASO {data.get('case_number', 'S/N')}: {subject}")
        elif kind == "soporte":
            if original == "Change":
                change_boxes.append(box)
            elif original == "TxDxSecure":
                _set_textbox_text(box, prepared_by, "000000")
            elif original == "Informe de Soporte":
                _set_textbox_text(box, title)

    if kind == "soporte":
        sequence_fallback = [report_date, report_date, client, client, subject, subject]
        for index, box in enumerate(change_boxes):
            styles = _textbox_styles(box)
            value = sequence_fallback[index] if index < len(sequence_fallback) else subject
            if "margin-left:214.15pt" in styles:
                value = report_date
            elif "margin-left:-57.8pt" in styles:
                value = client
            elif "margin-left:-33.9pt" in styles:
                value = subject
            _set_textbox_text(box, value, "000000")


def support_cover_dynamic_values(data: dict[str, Any]) -> list[str]:
    return [
        str(data.get("client_name") or data.get("establishment") or "Cliente")[:55],
        str(data.get("subject") or "Informe de soporte")[:75],
        _report_date_spanish(data),
        str(data.get("prepared_by") or "TxDxSecure")[:40],
    ]


def template_footer_text(document: Document) -> str:
    """Obtiene el pie visible de la plantilla antes de retirar sus secciones de ejemplo."""
    for section in document.sections:
        text = " ".join(paragraph.text.strip() for paragraph in section.footer.paragraphs if paragraph.text.strip())
        if text:
            return text
    return ""


def _footer_parts(kind: str, data: dict[str, Any], footer_text: str) -> tuple[str, str, str]:
    client = str(data.get("client_name") or data.get("establishment") or "Cliente")
    if kind == "small":
        return (
            "TXDXSECURE",
            f"Small Report {data.get('case_number', '')}",
            client,
        )
    center = f"Informe de Soporte - {client}"
    right = "Telf. 999 379 845 - 997 516 432"
    if footer_text and "\t" in footer_text:
        pieces = [piece.strip() for piece in footer_text.replace("Small Informe", center).split("\t")]
        if len(pieces) >= 3:
            return pieces[0], pieces[1], pieces[-1]
    return "www.txdxsecure.com", center, right


def restore_footer(document: Document, kind: str, data: dict[str, Any], footer_text: str) -> None:
    left, center, right = _footer_parts(kind, data, footer_text)
    for section in document.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.tab_stops.clear_all()
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(2.85))
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.15))
        for index, text in enumerate((left, center, right)):
            if index:
                paragraph.add_run("\t")
            run = paragraph.add_run(text)
            run.font.size = Pt(9)
            run.font.color.rgb = FOOTER_GRAY


def _style_name(document: Document, *names: str) -> str:
    for name in names:
        if name in document.styles:
            return name
    return names[-1]


def _format_run(
    run: Any,
    *,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    size: float | None = None,
) -> None:
    run.bold = bold or None
    run.italic = italic or None
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)


def _add_spacing(
    paragraph: Any,
    *,
    before: float = 0,
    after: float = 7,
    left: float = 0,
    first: float | None = None,
) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.left_indent = Pt(left) if left else None
    if first is not None:
        paragraph.paragraph_format.first_line_indent = Pt(first)


def _set_paragraph_bottom_border(paragraph: Any, color: str = TABLE_HEADER, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_section_heading(document: Document, text: str, number: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = _style_name(document, "Estilo1", "Heading 1")
    _add_spacing(paragraph, before=16, after=10)
    if number:
        number_run = paragraph.add_run(f"{number} ")
        _format_run(number_run, bold=True, color=BRAND_BLUE, size=13)
    run = paragraph.add_run(text)
    _format_run(run, bold=True, color=RGBColor(0x16, 0x34, 0x50), size=13)
    _set_paragraph_bottom_border(paragraph, TABLE_HEADER, "10")


def add_subheading(document: Document, text: str, number: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = _style_name(document, "Estilo2", "Heading 2")
    _add_spacing(paragraph, before=7, after=6, left=3)
    if number:
        number_run = paragraph.add_run(f"{number} ")
        _format_run(number_run, bold=True, color=BRAND_BLUE, size=11)
    run = paragraph.add_run(text)
    _format_run(run, bold=True, color=TEXT_GRAY, size=11)


def add_body_text(document: Document, text: Any, *, indent: bool = False, bold_label: str | None = None) -> None:
    chunks = [chunk.strip() for chunk in str(text or "No especificado").splitlines() if chunk.strip()] or ["No especificado"]
    for chunk in chunks:
        paragraph = document.add_paragraph()
        paragraph.style = _style_name(document, "Estilo3", "Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_spacing(paragraph, after=7, left=18 if indent else 0)
        if bold_label:
            run = paragraph.add_run(bold_label)
            _format_run(run, bold=True)
            paragraph.add_run(" ")
        run = paragraph.add_run(chunk)
        run.font.size = Pt(10)


def _split_text_items(value: Any) -> list[str]:
    if isinstance(value, list):
        return _as_list(value)
    text = str(value or "").strip()
    if not text:
        return []
    pieces = [piece.strip(" -\t") for piece in re.split(r"\n+|;\s+", text) if piece.strip(" -\t")]
    return pieces or [text]


def add_bullets(document: Document, values: Any) -> None:
    for value in _split_text_items(values):
        paragraph = document.add_paragraph()
        paragraph.style = _style_name(document, "Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_spacing(paragraph, after=4, left=23.1, first=-12.9)
        bullet = paragraph.add_run("• ")
        _format_run(bullet, bold=True, color=BRAND_BLUE, size=10)
        run = paragraph.add_run(value)
        run.font.size = Pt(10)


def add_numbered_items(document: Document, values: Any) -> None:
    for index, value in enumerate(_split_text_items(values), start=1):
        paragraph = document.add_paragraph()
        paragraph.style = _style_name(document, "Estilo4", "Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_spacing(paragraph, after=5, left=28, first=-14)
        label = paragraph.add_run(f"{index}. ")
        _format_run(label, bold=True, color=BRAND_BLUE)
        paragraph.add_run(value)


def add_code_block(document: Document, values: Any) -> None:
    for value in _split_text_items(values):
        paragraph = document.add_paragraph()
        paragraph.style = _style_name(document, "Normal")
        _add_spacing(paragraph, after=1, left=36)
        run = paragraph.add_run(value)
        run.font.name = "Consolas"
        _format_run(run, bold=True, size=9.5)
    document.add_paragraph()


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = _style_name(document, "Estilo7", "Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_spacing(paragraph, before=2, after=8)
    run = paragraph.add_run(text)
    _format_run(run, italic=True, color=TEXT_GRAY, size=8.5)


def add_note_label(document: Document, text: str = "Observación:") -> None:
    paragraph = document.add_paragraph()
    paragraph.style = _style_name(document, "Estilo5", "Normal")
    _add_spacing(paragraph, before=4, after=4)
    run = paragraph.add_run(text)
    _format_run(run, bold=True, color=TEXT_GRAY)


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_border(cell: Any, color: str = BORDER_LIGHT, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def _set_cell_margins(cell: Any, top: int = 120, start: int = 160, bottom: int = 120, end: int = 160) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_text(cell: Any, text: Any, *, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    _add_spacing(paragraph, after=2)
    run = paragraph.add_run(str(text or "No especificado"))
    _format_run(run, bold=bold, color=color, size=9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_callout_box(
    document: Document,
    title: str,
    content: Any = "",
    *,
    bullets: Any = None,
    fill: str = CALLOUT_BLUE,
    border: str = BORDER_LIGHT,
    title_color: RGBColor | None = None,
) -> None:
    text = _meaningful_text(content)
    bullet_items = [_meaningful_text(item) for item in _split_text_items(bullets)] if bullets is not None else []
    bullet_items = [item for item in bullet_items if item]
    if not text and not bullet_items:
        return

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    _shade_cell(cell, fill)
    _set_cell_border(cell, border)
    _set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    paragraph = cell.paragraphs[0]
    _add_spacing(paragraph, after=5)
    run = paragraph.add_run(title)
    _format_run(run, bold=True, color=title_color or BRAND_BLUE, size=10.5)

    if text:
        for chunk in [chunk.strip() for chunk in text.splitlines() if chunk.strip()]:
            body = cell.add_paragraph()
            body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_spacing(body, after=4)
            run = body.add_run(chunk)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x24, 0x32, 0x3D)

    for item in bullet_items:
        bullet = cell.add_paragraph()
        _add_spacing(bullet, after=3, left=14, first=-9)
        mark = bullet.add_run("• ")
        _format_run(mark, bold=True, color=title_color or BRAND_BLUE, size=9.5)
        run = bullet.add_run(item)
        run.font.size = Pt(9.5)
    document.add_paragraph()


def add_case_overview(document: Document, kind: str, data: dict[str, Any]) -> None:
    values = [
        ("Cliente", data.get("client_name") or data.get("establishment") or "Cliente"),
        ("Caso", data.get("case_number") or "S/N"),
        ("Fecha", _report_date_spanish(data)),
        ("Tipo", "Small Report" if kind == "small" else "Informe de Soporte"),
    ]
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cells = table.rows[0].cells
    for cell, (label, value) in zip(cells, values):
        _shade_cell(cell, CALLOUT_GRAY)
        _set_cell_border(cell, BORDER_LIGHT)
        _set_cell_margins(cell, top=90, start=100, bottom=90, end=100)
        label_paragraph = cell.paragraphs[0]
        _add_spacing(label_paragraph, after=2)
        label_run = label_paragraph.add_run(label.upper())
        _format_run(label_run, bold=True, color=BRAND_BLUE, size=7.5)
        value_paragraph = cell.add_paragraph()
        _add_spacing(value_paragraph, after=0)
        value_run = value_paragraph.add_run(str(value or "No especificado"))
        _format_run(value_run, bold=True, color=RGBColor(0x16, 0x34, 0x50), size=8.5)
    document.add_paragraph()


def add_key_values(document: Document, values: list[tuple[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    for cell, title in zip(header, ("Parámetro", "Detalle")):
        _shade_cell(cell, TABLE_HEADER)
        _set_cell_text(cell, title, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for index, (key, value) in enumerate(values, start=1):
        row = table.add_row().cells
        if index % 2 == 0:
            for cell in row:
                _shade_cell(cell, TABLE_ALT)
        _set_cell_text(row[0], key, bold=True, color=TEXT_GRAY)
        _set_cell_text(row[1], value)
    document.add_paragraph()


def build_small(document: Document, data: dict[str, Any], sections: dict[str, str]) -> None:
    add_section_heading(document, "Descripción del Soporte", "1.")
    add_subheading(document, f"Caso {data.get('case_number', 'S/N')}", "1.1.")
    add_body_text(document, data.get("subject") or data.get("request_description"), indent=True)
    add_subheading(document, "Cliente", "1.2.")
    add_body_text(document, data.get("client_name") or data.get("establishment"), indent=True)
    add_subheading(document, "Fecha", "1.3.")
    add_body_text(document, _report_date_spanish(data), indent=True)
    add_subheading(document, "Forma de Conexión de Soporte", "1.4.")
    add_body_text(document, data.get("connection_method"), indent=True)
    add_subheading(document, "Datos Base", "1.5.")
    add_body_text(document, data.get("request_description"), indent=True)

    add_section_heading(document, "Pasos realizados", "2.")
    add_subheading(document, "Relevamiento de Información", "2.1.")
    add_body_text(document, sections["information_gathering"])
    add_note_label(document)
    add_body_text(document, "La información se documenta en base a los datos confirmados para el tenant y el caso solicitado.")

    add_subheading(document, "Investigación de posibles causas", "2.2.")
    add_body_text(document, "Se consideraron las siguientes hipótesis operativas:")
    add_bullets(document, data.get("possible_causes") or sections["possible_causes"])

    add_subheading(document, "Secuencia de validación", "2.3.")
    add_numbered_items(document, data.get("actions_performed") or sections["actions"])
    if data.get("commands"):
        add_subheading(document, "Comandos", "2.4.")
        add_code_block(document, data.get("commands"))

    add_section_heading(document, "Acciones Realizadas", "3.")
    add_bullets(document, data.get("actions_performed") or sections["actions"])

    add_section_heading(document, "Resultados", "4.")
    add_bullets(document, data.get("results") or sections["results"])

    add_section_heading(document, "Recomendaciones", "5.")
    add_bullets(document, data.get("recommendations") or sections["recommendations"])


def build_support(document: Document, data: dict[str, Any], sections: dict[str, str]) -> None:
    add_section_heading(document, "Descripción del Soporte", "1.")
    add_subheading(document, "Tipo de Soporte", "1.1.")
    add_body_text(document, data.get("support_type"), indent=True)
    add_subheading(document, "Forma de Conexión de Soporte", "1.2.")
    add_body_text(document, data.get("connection_method"), indent=True)
    add_subheading(document, "Fecha y tiempo del soporte", "1.3.")
    add_body_text(document, f"{_report_date_spanish(data)}, {data.get('support_duration', 'duración no especificada')}", indent=True)
    add_subheading(document, "Nombre del establecimiento", "1.4.")
    add_body_text(document, data.get("establishment") or data.get("client_name"), indent=True)
    add_subheading(document, "Objetivo", "1.5.")
    add_body_text(document, data.get("objective") or data.get("subject"), indent=True)
    add_subheading(document, "Ubicación", "1.6.")
    add_body_text(document, data.get("location"), indent=True)

    if data.get("technical_details"):
        add_subheading(document, "Datos técnicos", "1.7.")
        add_key_values(document, list(dict(data["technical_details"]).items()))
        add_caption(document, "Tabla 1. Datos técnicos asociados al caso de soporte.")

    add_section_heading(document, "Antecedentes y Escenario Inicial", "2.")
    add_body_text(document, sections["background"])
    add_body_text(document, sections["initial_scenario"])
    add_note_label(document, "Evaluación de impacto:")
    add_bullets(document, data.get("impact_notes") or [
        "La actividad se documenta con los datos confirmados del tenant solicitado.",
        "Las validaciones se registran sin asumir evidencia no entregada en la solicitud.",
    ])

    add_section_heading(document, "Desarrollo de la Actividad", "3.")
    add_subheading(document, "Acciones Realizadas", "3.1.")
    add_body_text(document, sections["actions"])
    add_bullets(document, data.get("actions_performed") or sections["actions"])
    if data.get("commands"):
        add_subheading(document, "Comandos ejecutados", "3.2.")
        add_code_block(document, data.get("commands"))

    add_section_heading(document, "Resultados Obtenidos", "4.")
    add_body_text(document, sections["results"])
    add_bullets(document, data.get("results") or sections["results"])
    validations = data.get("validations") or []
    if validations:
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell, title in zip(table.rows[0].cells, ("Validación", "Resultado", "Estado")):
            _shade_cell(cell, TABLE_HEADER)
            _set_cell_text(cell, title, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        for index, item in enumerate(validations, start=1):
            row = table.add_row().cells
            if index % 2 == 0:
                for cell in row:
                    _shade_cell(cell, TABLE_ALT)
            _set_cell_text(row[0], item.get("validation", ""))
            _set_cell_text(row[1], item.get("result", ""))
            _set_cell_text(row[2], item.get("status", ""), bold=True, color=BRAND_BLUE)
        add_caption(document, "Tabla 2. Validaciones posteriores a la actividad.")

    add_section_heading(document, "Recomendaciones", "5.")
    add_bullets(document, data.get("recommendations") or sections["recommendations"])

    paragraph = document.add_paragraph()
    paragraph.style = _style_name(document, "Normal")
    _add_spacing(paragraph, after=7)
    run = paragraph.add_run("Conclusión:")
    _format_run(run, bold=True)
    add_body_text(document, sections["conclusion"])


EMPTY_MEANING_PHRASES = {
    "no especificado",
    "no especificada",
    "sin informacion",
    "sin informacion adicional",
    "sin información",
    "sin información adicional",
    "sin datos",
    "sin datos confirmados",
    "sin acciones confirmadas",
    "sin resultados confirmados",
    "sin recomendaciones confirmadas",
    "no determinado",
    "no determinada",
    "pendiente",
}


def _meaningful_text(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    normalized = re.sub(r"\s+", " ", text.lower()).strip(" .:-")
    if normalized in EMPTY_MEANING_PHRASES:
        return ""
    return text


def _dynamic_section(item: Any) -> dict[str, Any] | None:
    if not isinstance(item, dict):
        return None
    title = _meaningful_text(item.get("title"))
    content = _meaningful_text(item.get("content"))
    bullets = [_meaningful_text(value) for value in item.get("bullets", [])] if isinstance(item.get("bullets"), list) else []
    bullets = [value for value in bullets if value]
    subsections = [_dynamic_section(value) for value in item.get("subsections", [])] if isinstance(item.get("subsections"), list) else []
    subsections = [value for value in subsections if value]
    if not title or (not content and not bullets and not subsections):
        return None
    return {"title": title, "content": content, "bullets": bullets, "subsections": subsections}


def dynamic_sections_from_payload(payload: dict[str, Any], data: dict[str, Any], kind: str) -> list[dict[str, Any]]:
    raw_sections = payload.get("sections")
    if isinstance(raw_sections, list):
        sections = [_dynamic_section(item) for item in raw_sections]
        sections = [section for section in sections if section]
        if sections:
            return sections

    summary = _meaningful_text(payload.get("summary") or data.get("request_description") or data.get("background"))
    if summary:
        return [{"title": "Descripcion del Soporte", "content": summary, "bullets": [], "subsections": []}]

    legacy_pairs = (
        [
            ("Descripcion del Soporte", payload.get("technical_evidence")),
            ("Observaciones", payload.get("observations")),
            ("Recomendaciones", payload.get("recommended_actions")),
            ("Conclusion", payload.get("conclusion")),
        ]
        if kind == "small"
        else [
            ("Descripcion del Soporte", payload.get("case_context")),
            ("Antecedentes y Escenario Inicial", payload.get("visual_evidence_summary")),
            ("Analisis Tecnico", payload.get("technical_analysis")),
            ("Acciones Realizadas", payload.get("actions_performed")),
            ("Resultados Obtenidos", payload.get("results_obtained")),
            ("Recomendaciones", payload.get("recommended_actions")),
            ("Conclusion", payload.get("conclusion")),
        ]
    )
    sections: list[dict[str, Any]] = []
    for title, value in legacy_pairs:
        items = _as_list(value)
        content = _meaningful_text(value if not isinstance(value, list) else "")
        bullets = [_meaningful_text(item) for item in items] if isinstance(value, list) else []
        bullets = [item for item in bullets if item]
        if content or bullets:
            sections.append({"title": title, "content": content, "bullets": bullets, "subsections": []})
    return sections


def build_dynamic_report(document: Document, kind: str, payload: dict[str, Any], data: dict[str, Any]) -> None:
    """Construye el cuerpo con secciones condicionales generadas por Azure."""
    sections = dynamic_sections_from_payload(payload, data, kind)
    document.add_page_break()
    add_case_overview(document, kind, data)
    summary = _meaningful_text(payload.get("summary"))
    if summary:
        add_callout_box(document, "Resumen ejecutivo", summary, fill=CALLOUT_BLUE, border=TABLE_HEADER)
    if sections:
        intro = document.add_paragraph()
        intro.style = _style_name(document, "Estilo1", "Heading 1")
        _add_spacing(intro, before=2, after=10)
        run = intro.add_run("Contenido del reporte")
        _format_run(run, bold=True, color=BRAND_BLUE, size=12)
        _set_paragraph_bottom_border(intro, BORDER_LIGHT, "6")
    else:
        add_callout_box(
            document,
            "Sin secciones confirmadas",
            "Azure no devolvio secciones sustentadas por la evidencia entregada.",
            fill=CALLOUT_YELLOW,
            border="E2B93B",
            title_color=RGBColor(0x8A, 0x61, 0x00),
        )
    section_number = 1
    for section in sections:
        add_section_heading(document, section["title"], f"{section_number}.")
        if section.get("content"):
            add_body_text(document, section["content"])
        if section.get("bullets"):
            add_bullets(document, section["bullets"])
        for subsection_number, subsection in enumerate(section.get("subsections", []), start=1):
            add_subheading(document, subsection["title"], f"{section_number}.{subsection_number}.")
            if subsection.get("content"):
                add_body_text(document, subsection["content"], indent=True)
            if subsection.get("bullets"):
                add_bullets(document, subsection["bullets"])
        section_number += 1

    limitations = [_meaningful_text(item) for item in _as_list(payload.get("limitations"))]
    limitations = [item for item in limitations if item]
    if limitations:
        add_callout_box(
            document,
            f"{section_number}. Limitaciones",
            bullets=limitations,
            fill=CALLOUT_YELLOW,
            border="E2B93B",
            title_color=RGBColor(0x8A, 0x61, 0x00),
        )


def validate_docx(path: Path) -> None:
    if not path.exists() or path.stat().st_size < 10_000:
        raise RuntimeError("El DOCX no fue generado correctamente")
    with zipfile.ZipFile(path) as archive:
        if "word/document.xml" not in archive.namelist():
            raise RuntimeError("El DOCX generado no contiene document.xml")


def main() -> None:
    parser = argparse.ArgumentParser(description="Genera Small Report o Informe de Soporte desde una plantilla XOC.")
    parser.add_argument("type", choices=("small", "soporte"), help="small para resumen corto; soporte para informe técnico.")
    parser.add_argument("data", type=Path, help="Ruta al JSON con los datos confirmados de la atención.")
    parser.add_argument("--no-azure", action="store_true", help="Solo para prueba local: no redacta secciones mediante Azure.")
    parser.add_argument("--allow-local-fallback", action="store_true", help="Si Azure falla, genera el DOCX con los textos de entrada.")
    args = parser.parse_args()
    load_dotenv(ROOT / ".env")

    if not args.data.is_file():
        raise SystemExit(f"No existe el archivo de datos: {args.data}")
    template = TEMPLATES[args.type]
    if not template.is_file():
        raise SystemExit(f"No existe la plantilla: {template}")
    data = json.loads(args.data.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise SystemExit("El archivo de datos debe contener un objeto JSON")

    payload = data if isinstance(data.get("sections"), list) else {
        "title": str(data.get("report_title") or data.get("subject") or "Reporte XOC"),
        "summary": str(data.get("request_description") or data.get("background") or data.get("subject") or ""),
        "sections": [],
        "limitations": _as_list(data.get("limitations")),
        "image_citations": [],
    }

    document = Document(template)
    footer_text = template_footer_text(document)
    update_cover_fields(document, args.type, data)
    clear_document_body_except_cover(document)
    restore_footer(document, args.type, data, footer_text)
    build_dynamic_report(document, args.type, payload, data)
    OUTPUT_DIR.mkdir(exist_ok=True)
    result = output_path(args.type, data)
    document.save(result)
    if args.type != "small":
        patch_support_cover_text_color(result, support_cover_dynamic_values(data))
    validate_docx(result)
    print(f"Reporte generado: {result}")


if __name__ == "__main__":
    main()
