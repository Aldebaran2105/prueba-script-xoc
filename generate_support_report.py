"""Genera Informe de Soporte desde texto + imagenes usando Foundry multimodal."""

from __future__ import annotations

import argparse
import json
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv

from azure_foundry_multimodal_client import DEFAULT_PAYLOAD_PATH, generate_report_payload
from generate_report import (
    TEMPLATES,
    OUTPUT_DIR,
    build_dynamic_report,
    clear_document_body_except_cover,
    output_path,
    patch_support_cover_text_color,
    restore_footer,
    support_cover_dynamic_values,
    template_footer_text,
    update_cover_fields,
    validate_docx,
    add_caption,
    add_section_heading,
)

ROOT = Path(__file__).resolve().parent


def _load_json(path: Path | None) -> dict[str, Any]:
    if not path:
        return {}
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise SystemExit("El JSON estructurado debe contener un objeto.")
    return data


def _text_from_args(args: argparse.Namespace) -> str:
    pieces = []
    if args.text:
        pieces.append(args.text)
    if args.text_file:
        pieces.append(args.text_file.read_text(encoding="utf-8"))
    return "\n\n".join(piece.strip() for piece in pieces if piece.strip())


def _as_list(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if value:
        return [str(value).strip()]
    return []


def _payload_to_data(payload: dict[str, Any], structured: dict[str, Any], text: str) -> dict[str, Any]:
    technical_details = dict(structured.get("technical_details") or {})
    return {
        **structured,
        "case_number": structured.get("case_number") or "IF-MULTIMODAL",
        "client_name": structured.get("client_name") or structured.get("establishment") or "Cliente",
        "report_date": structured.get("report_date") or date.today().isoformat(),
        "support_type": structured.get("support_type") or "Soporte tecnico",
        "connection_method": structured.get("connection_method") or "No especificado",
        "support_duration": structured.get("support_duration") or "No especificada",
        "establishment": structured.get("establishment") or structured.get("client_name") or "Cliente",
        "location": structured.get("location") or "No especificada",
        "subject": structured.get("subject") or payload.get("title") or "Informe de Soporte",
        "objective": structured.get("objective") or payload.get("title") or "Documentar evidencia tecnica del caso.",
        "background": structured.get("background") or payload.get("summary") or text,
        "technical_details": technical_details,
        "validations": structured.get("validations") or [],
    }


def _sections_from_payload(payload: dict[str, Any], data: dict[str, Any]) -> dict[str, str]:
    technical_analysis = str(payload.get("technical_analysis") or "Sin analisis tecnico confirmado.")
    visual_summary = str(payload.get("visual_evidence_summary") or "Sin evidencia visual confirmada.")
    return {
        "background": str(data.get("background") or "Sin antecedentes confirmados."),
        "initial_scenario": f"{visual_summary}\n\n{technical_analysis}",
        "actions": "; ".join(_as_list(data.get("actions_performed"))) or "Sin acciones confirmadas.",
        "results": "; ".join(_as_list(data.get("results"))) or str(payload.get("results_obtained") or "Sin resultados confirmados."),
        "recommendations": "; ".join(_as_list(data.get("recommendations"))) or "Sin recomendaciones confirmadas.",
        "conclusion": str(payload.get("conclusion") or "Sin conclusion confirmada."),
    }


def _append_evidence_images(document: Document, image_paths: list[Path]) -> None:
    if not image_paths:
        return
    add_section_heading(document, "Evidencia Visual", "6.")
    for index, path in enumerate(image_paths, start=1):
        paragraph = document.add_paragraph()
        paragraph.alignment = 1
        paragraph.add_run().add_picture(str(path), width=Inches(5.8))
        add_caption(document, f"Figura {index}. Evidencia visual proporcionada para el informe de soporte.")


def _writable_output_path(path: Path) -> Path:
    if not path.exists():
        return path
    try:
        with path.open("r+b"):
            return path
    except PermissionError:
        for version in range(2, 100):
            candidate = path.with_name(f"{path.stem}_v{version}{path.suffix}")
            if not candidate.exists():
                return candidate
        raise


def main() -> None:
    parser = argparse.ArgumentParser(description="Genera Informe de Soporte desde texto e imagenes.")
    parser.add_argument("--text", default="", help="Texto libre del analista.")
    parser.add_argument("--text-file", type=Path, help="Archivo TXT/MD con notas del analista.")
    parser.add_argument("--data", type=Path, help="JSON opcional con datos estructurados.")
    parser.add_argument("--image", action="append", default=[], type=Path, help="Ruta de imagen local. Se puede repetir.")
    parser.add_argument("--no-azure", action="store_true", help="Genera payload local sin llamar a Azure.")
    parser.add_argument("--allow-local-fallback", action="store_true", default=True, help="Si Azure falla, genera borrador local.")
    args = parser.parse_args()
    load_dotenv(ROOT / ".env")

    text = _text_from_args(args)
    structured = _load_json(args.data)
    payload = generate_report_payload(
        "informe_soporte",
        text=text,
        image_paths=args.image,
        structured_data=structured,
        use_azure=not args.no_azure,
        allow_local_fallback=args.allow_local_fallback,
        output_path=DEFAULT_PAYLOAD_PATH,
    )
    data = _payload_to_data(payload, structured, text)
    document = Document(TEMPLATES["soporte"])
    footer_text = template_footer_text(document)
    update_cover_fields(document, "soporte", data)
    clear_document_body_except_cover(document)
    restore_footer(document, "soporte", data, footer_text)
    build_dynamic_report(document, "soporte", payload, data)
    _append_evidence_images(document, args.image)

    OUTPUT_DIR.mkdir(exist_ok=True)
    result = _writable_output_path(output_path("soporte", data))
    document.save(result)
    patch_support_cover_text_color(result, support_cover_dynamic_values(data))
    validate_docx(result)
    print(f"Payload JSON: {DEFAULT_PAYLOAD_PATH}")
    print(f"Informe de Soporte generado: {result}")


if __name__ == "__main__":
    main()
